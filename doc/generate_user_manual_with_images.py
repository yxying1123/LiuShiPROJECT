#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
单细胞数据分析系统 - 用户功能手册生成器（带截图版）
生成docx格式的分模块用户手册，包含系统界面截图
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', 'Microsoft YaHei', 'SimSun']
    font_size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12)
    set_chinese_font(run, font_name='SimHei', font_size=font_size, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph_zh(doc, text, bold=False, font_size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加中文段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_chinese_font(run, bold=bold, font_size=font_size)
    return para


def add_image_with_caption(doc, image_path, caption, width=6.0):
    """添加图片及说明文字"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))

        # 添加图片说明
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(f'图：{caption}')
        set_chinese_font(caption_run, font_size=9, bold=False)
        caption_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # 添加空行
        return True
    return False


def add_table_zh(doc, headers, rows):
    """添加中文表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, bold=True, font_size=10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_size=10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_code_block(doc, code_text):
    """添加代码块样式段落"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def add_note_box(doc, title, content):
    """添加提示框"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # 标题
    run_title = para.add_run(f"【{title}】")
    set_chinese_font(run_title, bold=True, font_size=10)
    run_title.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    # 内容
    run_content = para.add_run(content)
    set_chinese_font(run_content, font_size=10)


def add_step_list(doc, steps):
    """添加步骤列表"""
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(step)
        set_chinese_font(run, font_size=10.5)


def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(item)
        set_chinese_font(run, font_size=10.5)


def create_manual():
    """创建用户手册"""
    doc = Document()
    pic_dir = '/Users/yangxiaoyu/学习/all/all-projectR/doc/pic'

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # ========== 封面 ==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(150)
    title_run = title_para.add_run('单细胞数据分析系统')
    set_chinese_font(title_run, font_name='SimHei', font_size=28, bold=True)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(30)
    subtitle_run = subtitle_para.add_run('用户功能手册')
    set_chinese_font(subtitle_run, font_name='SimHei', font_size=22, bold=True)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.paragraph_format.space_before = Pt(100)
    version_run = version_para.add_run('版本：v1.0.0\n日期：2024年1月')
    set_chinese_font(version_run, font_size=12)

    doc.add_page_break()

    # ========== 目录 ==========
    add_heading_zh(doc, '目录', level=1)
    toc_items = [
        '1. 系统概述',
        '2. 文件管理模块',
        '3. 散点图分析模块',
        '4. 聚类分析模块',
        '5. 热图聚类模块',
        '6. 数据导出模块',
        '7. 常见问题与解决方案',
        '8. 附录'
    ]
    for item in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        run = para.add_run(item)
        set_chinese_font(run, font_size=11)

    doc.add_page_break()

    # ========== 第一章：系统概述 ==========
    add_heading_zh(doc, '1. 系统概述', level=1)

    add_heading_zh(doc, '1.1 产品简介', level=2)
    add_paragraph_zh(doc,
        '单细胞数据分析系统是一款面向流式细胞术（Flow Cytometry）和质谱流式（CyTOF）单细胞数据的专业分析平台。'
        '系统提供从数据上传、降维可视化、聚类分析到热图展示的一站式解决方案，帮助研究人员快速发现数据中的生物学模式和细胞亚群。')

    add_heading_zh(doc, '1.2 核心功能模块', level=2)
    add_paragraph_zh(doc, '系统包含以下核心功能模块：')

    headers = ['功能模块', '功能描述']
    rows = [
        ['文件管理模块', '支持 CSV、FCS 等多种格式文件的上传、存储和管理'],
        ['散点图分析模块', '基于 UMAP 算法的高维数据降维和散点图展示'],
        ['聚类分析模块', '基于 Phenograph 算法的细胞聚类，自动识别细胞亚群'],
        ['热图聚类模块', '层次聚类热图展示，支持行/列聚类树可视化'],
        ['数据导出模块', '支持分析结果的图片、CSV 数据导出']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '1.3 适用场景', level=2)
    add_paragraph_zh(doc, '本系统适用于以下研究场景：')
    scenarios = [
        '免疫细胞分群分析',
        '肿瘤微环境细胞异质性研究',
        '药物处理后细胞状态变化分析',
        '单细胞水平的生物标志物发现'
    ]
    add_bullet_list(doc, scenarios)

    add_heading_zh(doc, '1.4 系统架构', level=2)
    add_paragraph_zh(doc, '系统采用前后端分离架构：')
    add_paragraph_zh(doc, '前端（cluster）：基于 React + Vite 构建，提供文件管理、散点图分析、热图聚类分析等功能界面。')
    add_paragraph_zh(doc, '后端（cluster-End）：基于 FastAPI (Python) 构建，提供文件服务、UMAP 降维算法、Phenograph 聚类算法等核心计算服务。')

    doc.add_page_break()

    # ========== 第二章：文件管理模块 ==========
    add_heading_zh(doc, '2. 文件管理模块', level=1)

    add_heading_zh(doc, '2.1 模块概述', level=2)
    add_paragraph_zh(doc,
        '文件管理模块是系统的数据入口，提供数据文件的上传、查看、删除和下载功能。'
        '支持批量上传多个文件，自动识别文件格式和列信息。')

    # 插入文件列表面面截图
    add_image_with_caption(doc, os.path.join(pic_dir, '01_file_list.png'),
                          '文件列表页面 - 展示已上传的数据文件', width=6.5)

    add_heading_zh(doc, '2.2 支持的文件格式', level=2)
    headers = ['格式', '扩展名', '说明']
    rows = [
        ['CSV', '.csv', '逗号分隔值文件，最常用格式'],
        ['文本', '.txt', '制表符或空格分隔的数据文件'],
        ['FCS', '.fcs', '流式细胞术标准数据格式']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '2.3 功能操作说明', level=2)

    add_heading_zh(doc, '2.3.1 上传文件', level=3)
    add_paragraph_zh(doc, '操作步骤：')
    upload_steps = [
        '点击左侧导航栏"文件列表"进入文件管理页面',
        '点击"FCS文件导入"或"CSV文件导入"按钮，或直接将文件拖拽到上传区域',
        '在弹出的对话框中选择一个或多个文件（支持批量上传）',
        '系统自动验证文件格式，上传完成后显示文件列表',
        '点击文件名可查看列信息和数值列详情'
    ]
    add_step_list(doc, upload_steps)

    add_note_box(doc, '提示', '支持同时上传多个文件，系统会自动检测文件格式并提取列信息。')

    add_heading_zh(doc, '2.3.2 查看文件信息', level=3)
    add_paragraph_zh(doc, '操作步骤：')
    view_steps = [
        '在文件列表中点击目标文件名',
        '系统显示文件详细信息，包括所有列名',
        '数值列会被自动识别并标注',
        '确认数据格式正确后即可用于后续分析'
    ]
    add_step_list(doc, view_steps)

    add_heading_zh(doc, '2.3.3 删除文件', level=3)
    add_paragraph_zh(doc, '操作步骤：')
    delete_steps = [
        '在文件列表中找到要删除的文件',
        '点击文件卡片上的删除按钮',
        '在确认对话框中点击"确认"完成删除'
    ]
    add_step_list(doc, delete_steps)

    add_heading_zh(doc, '2.3.4 下载文件', level=3)
    add_paragraph_zh(doc, '操作步骤：')
    download_steps = [
        '在文件列表中找到要下载的文件',
        '点击下载按钮即可将文件保存到本地'
    ]
    add_step_list(doc, download_steps)

    add_heading_zh(doc, '2.4 数据格式要求', level=2)
    add_paragraph_zh(doc, 'CSV 文件格式要求：')
    csv_requirements = [
        '第一行为列名（header）',
        '使用逗号作为分隔符',
        '数值列应为数字格式',
        '支持缺失值（留空或填写 NA）',
        '文件编码推荐 UTF-8'
    ]
    add_bullet_list(doc, csv_requirements)

    add_paragraph_zh(doc, 'CSV 文件示例：')
    csv_example = '''Time,Event_length,CD3,CD4,CD8,Sample
1,12,150.5,200.3,50.2,Sample1
2,15,180.2,220.1,45.8,Sample1
3,11,120.8,180.5,60.3,Sample2'''
    add_code_block(doc, csv_example)

    doc.add_page_break()

    # ========== 第三章：散点图分析模块 ==========
    add_heading_zh(doc, '3. 散点图分析模块', level=1)

    add_heading_zh(doc, '3.1 模块概述', level=2)
    add_paragraph_zh(doc,
        '散点图分析模块是系统的核心功能之一，提供高维单细胞数据的 UMAP 降维可视化。'
        '支持多种数据选择模式、降维重计算和交互式框选功能。')

    # 插入数据分析页面截图
    add_image_with_caption(doc, os.path.join(pic_dir, '02_scatter_analysis.png'),
                          '数据分析页面 - 散点图分析主界面', width=6.5)

    add_heading_zh(doc, '3.2 工作流程', level=2)
    workflow_steps = [
        '选择文件：从已上传的文件中选择要分析的数据',
        '选择变量：选择 X 轴和 Y 轴变量，或选择多维度进行 UMAP 降维',
        '生成散点图：系统计算并展示散点图',
        '交互选择：使用框选工具选择感兴趣的数据区域',
        '聚类分析：对选中的数据进行聚类分析'
    ]
    for i, step in enumerate(workflow_steps, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'步骤{i}：{step}')
        set_chinese_font(run, font_size=10.5)

    add_heading_zh(doc, '3.3 分析模式', level=2)

    add_heading_zh(doc, '3.3.1 初始配置对话框', level=3)
    add_paragraph_zh(doc, '点击"新建分析"按钮后，会弹出初始配置对话框：')

    # 插入初始配置截图
    add_image_with_caption(doc, os.path.join(pic_dir, '03_initial_config.png'),
                          '初始配置对话框 - 选择分析模式、文件和维度', width=5.5)

    add_paragraph_zh(doc, '配置选项说明：')
    config_items = [
        '分析模式：选择"降维散点图"（UMAP降维）或"二维散点图"（直接使用已有坐标）',
        '选择文件：勾选要分析的数据文件，支持多文件联合分析',
        '维度选择：选择参与降维的标记物/特征列',
        '随机读取行数：设置采样行数，留空表示读取全部数据'
    ]
    add_bullet_list(doc, config_items)

    add_heading_zh(doc, '3.3.2 维度选择', level=3)
    add_paragraph_zh(doc, '点击"维度选择"下拉框，可以选择要参与分析的变量：')

    # 插入维度选择截图
    add_image_with_caption(doc, os.path.join(pic_dir, '04_dimension_select.png'),
                          '维度选择界面 - 勾选需要分析的标记物', width=5.0)

    add_note_box(doc, '说明', 'FSC-A、SSC-A、FSC-H、FSC-W、Time 列为系统保留列，不可选择用于降维分析。')

    add_heading_zh(doc, '3.4 UMAP 降维参数', level=2)
    headers = ['参数', '默认值', '说明']
    rows = [
        ['n_neighbors', '15', '局部邻域大小，影响局部/全局结构平衡'],
        ['min_dist', '0.001', '嵌入空间中点的最小距离'],
        ['target_weight', '0.5', '目标权重'],
        ['metric', 'euclidean', '距离度量方式']
    ]
    add_table_zh(doc, headers, rows)

    add_note_box(doc, '参数调优建议',
        'n_neighbors：小值（5-15）保留更多局部结构，大值（50-100）保留更多全局结构；'
        'min_dist：小值使点更聚集，大值使点更分散。')

    add_heading_zh(doc, '3.5 交互功能', level=2)
    headers = ['功能', '操作方式', '说明']
    rows = [
        ['框选', '鼠标拖拽', '选择感兴趣的数据区域'],
        ['缩放', '滚轮/手势', '放大/缩小查看细节'],
        ['平移', '拖拽画布', '移动视图位置'],
        ['悬停', '鼠标悬停', '显示数据点详细信息'],
        ['图例筛选', '点击图例', '显示/隐藏特定来源数据']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '3.6 操作步骤详解', level=2)

    add_heading_zh(doc, '3.6.1 生成散点图', level=3)

    # 插入生成的散点图截图
    add_image_with_caption(doc, os.path.join(pic_dir, '05_scatter_generated.png'),
                          '生成的UMAP散点图 - 展示4000个细胞的降维分布', width=6.5)

    scatter_steps = [
        '进入散点图分析页面',
        '点击"新建分析"按钮',
        '在初始配置对话框中选择分析模式（推荐"降维散点图"）',
        '选择要分析的文件（支持多选）',
        '点击"维度选择"，勾选要分析的标记物',
        '设置采样行数（可选，大数据集建议设置）',
        '点击"确定"按钮',
        '等待数据处理完成，查看生成的散点图'
    ]
    add_step_list(doc, scatter_steps)

    add_heading_zh(doc, '3.6.2 框选数据', level=3)
    add_paragraph_zh(doc, '使用工具栏中的"Lasso Select"（套索选择）或"Box Select"（框选）工具：')
    select_steps = [
        '在散点图工具栏中选择选择工具',
        '在散点图上按住鼠标左键拖拽绘制选择区域',
        '释放鼠标完成框选',
        '可以多次框选累积选择多个区域',
        '选中的数据点会高亮显示，数量显示在"已选 X 个点"中'
    ]
    add_step_list(doc, select_steps)

    add_heading_zh(doc, '3.6.3 重新降维', level=3)
    add_paragraph_zh(doc, '框选感兴趣的数据区域后，可以点击"降维分析"按钮对选中数据进行重新降维。')

    doc.add_page_break()

    # ========== 第四章：聚类分析模块 ==========
    add_heading_zh(doc, '4. 聚类分析模块', level=1)

    add_heading_zh(doc, '4.1 模块概述', level=2)
    add_paragraph_zh(doc,
        '聚类分析模块基于 Leiden 算法对选中的数据点进行聚类分析，自动识别细胞亚群。'
        '聚类结果可用于后续的热图分析和 ROE（Relative Observed vs Expected）分析。')

    add_heading_zh(doc, '4.2 聚类算法', level=2)
    add_paragraph_zh(doc,
        '系统使用 Leiden 算法进行细胞聚类。该算法是一种基于图的聚类方法，'
        '能够高效地发现数据中的社区结构，广泛应用于单细胞数据分析。')

    add_heading_zh(doc, '4.3 算法参数', level=2)

    # 插入聚类参数对话框截图
    add_image_with_caption(doc, os.path.join(pic_dir, '06_cluster_dialog.png'),
                          '聚类分析参数对话框 - 设置聚类参数', width=5.0)

    headers = ['参数', '默认值', '说明']
    rows = [
        ['k (近邻数)', '20', 'k-近邻图中的邻居数，影响聚类粒度'],
        ['resolution', '1.0', '社区检测分辨率，影响聚类数量'],
        ['n_iterations', '20', '迭代次数'],
        ['seed', '123', '随机种子，保证结果可重复']
    ]
    add_table_zh(doc, headers, rows)

    add_note_box(doc, '参数说明',
        'k 值：小 k 产生更多小聚类，大 k 产生更少大聚类；'
        'resolution：分辨率参数，值越大产生的聚类数量越多。')

    add_heading_zh(doc, '4.4 操作流程', level=2)
    cluster_steps = [
        '在散点图页面生成散点图（无需框选，系统会对全部数据点进行聚类）',
        '点击左侧"聚类分析"按钮',
        '在弹出的"聚类分析参数"对话框中设置参数（可使用默认值）',
        '点击"开始聚类"按钮',
        '等待分析完成，系统自动跳转到结果聚类页面展示热图聚类树'
    ]
    add_step_list(doc, cluster_steps)

    add_heading_zh(doc, '4.5 结果展示', level=2)
    add_paragraph_zh(doc, '聚类分析完成后，系统会展示以下结果：')
    results = [
        '聚类热图树：展示各聚类的表达特征和层次关系',
        '聚类表：各聚类的均值矩阵，可导出为CSV',
        '散点表：包含聚类编号的原始数据表，可导出为CSV'
    ]
    add_bullet_list(doc, results)

    add_heading_zh(doc, '4.6 注意事项', level=2)
    notes = [
        '聚类分析需要一定计算时间，请耐心等待',
        '聚类结果受 k 值和 resolution 参数影响，可根据需要调整',
        '建议保存聚类结果以便后续查看和比较'
    ]
    add_bullet_list(doc, notes)

    doc.add_page_break()

    # ========== 第五章：热图聚类模块 ==========
    add_heading_zh(doc, '5. 热图聚类模块', level=1)

    add_heading_zh(doc, '5.1 模块概述', level=2)
    add_paragraph_zh(doc,
        '热图聚类模块展示聚类结果的表达矩阵，支持行和列的层次聚类，'
        '帮助识别基因/蛋白表达模式和样本间相似性。')

    add_heading_zh(doc, '5.2 视图选项', level=2)
    headers = ['视图', '说明']
    rows = [
        ['聚类热图树', '显示带层次聚类树的热图'],
        ['聚类表', '显示聚类均值矩阵数据表'],
        ['散点表', '显示原始散点数据表']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '5.3 聚类热图树', level=2)

    # 插入热图聚类树截图
    add_image_with_caption(doc, os.path.join(pic_dir, '07_heatmap_cluster.png'),
                          '聚类热图树 - 展示9个聚类的表达特征和层次聚类关系', width=6.5)

    add_paragraph_zh(doc, '热图聚类树包含以下元素：')
    heatmap_elements = [
        '行（左侧）：标记物/特征（如 CD3、CD4、CD8 等）',
        '列（上方）：聚类编号（如 0、1、2...）',
        '颜色：标准化后的表达值（红色表示高表达，蓝色表示低表达）',
        '树状图：行和列的层次聚类关系'
    ]
    add_bullet_list(doc, heatmap_elements)

    add_heading_zh(doc, '5.4 聚类表', level=2)

    # 插入聚类表截图
    add_image_with_caption(doc, os.path.join(pic_dir, '08_cluster_table.png'),
                          '聚类表 - 展示各聚类的均值矩阵数据', width=6.5)

    add_paragraph_zh(doc, '聚类表展示每个聚类中各标记物的平均表达值，可以：')
    table_features = [
        '查看各细胞亚群的表达特征',
        '比较不同聚类间的表达差异',
        '点击"下载聚类表"按钮导出为 CSV 文件'
    ]
    add_bullet_list(doc, table_features)

    add_heading_zh(doc, '5.5 操作步骤', level=2)

    add_heading_zh(doc, '5.5.1 查看聚类热图树', level=3)
    heatmap_steps = [
        '聚类分析完成后，系统会自动跳转到结果聚类页面',
        '默认显示"聚类热图树"标签页',
        '查看热图和层次聚类树',
        '观察表达模式和聚类关系',
        '点击"下载热图图片"按钮可保存为 PNG 图片'
    ]
    add_step_list(doc, heatmap_steps)

    add_heading_zh(doc, '5.5.2 查看聚类表', level=3)
    table_steps = [
        '切换到"聚类表"标签页',
        '查看各聚类的均值矩阵',
        '表格显示每个聚类中各标记物的平均表达值',
        '点击"下载聚类表"按钮可将数据导出为 CSV 文件'
    ]
    add_step_list(doc, table_steps)

    add_heading_zh(doc, '5.5.3 保存结果', level=3)
    save_steps = [
        '在结果聚类页面，点击右上角"保存结果"按钮',
        '输入结果名称',
        '结果将保存到分析结果页面，便于后续查看'
    ]
    add_step_list(doc, save_steps)

    doc.add_page_break()

    # ========== 第六章：数据导出模块 ==========
    add_heading_zh(doc, '6. 数据导出模块', level=1)

    add_heading_zh(doc, '6.1 模块概述', level=2)
    add_paragraph_zh(doc,
        '数据导出模块支持将分析结果导出为图片或 CSV 格式，便于后续分析和报告撰写。')

    add_heading_zh(doc, '6.2 导出选项', level=2)
    headers = ['导出内容', '格式', '说明']
    rows = [
        ['散点图', 'PNG', '高分辨率散点图图片'],
        ['热图', 'PNG', '热图聚类树图片'],
        ['聚类表', 'CSV', '聚类均值矩阵'],
        ['散点表', 'CSV', '包含聚类编号的原始数据表']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '6.3 导出操作', level=2)

    add_heading_zh(doc, '6.3.1 导出图片', level=3)
    img_export_steps = [
        '在散点图页面，使用工具栏中的"Download plot as a PNG"按钮下载散点图',
        '在热图聚类树页面，点击"下载热图图片"按钮',
        '选择保存位置，图片将以 PNG 格式保存到本地'
    ]
    add_step_list(doc, img_export_steps)

    add_heading_zh(doc, '6.3.2 导出数据表', level=3)
    data_export_steps = [
        '在聚类表页面，点击"下载聚类表"按钮',
        '在散点表页面，点击相应导出按钮',
        '数据将以 CSV 格式保存到本地'
    ]
    add_step_list(doc, data_export_steps)

    add_heading_zh(doc, '6.4 分析结果页面', level=2)

    # 插入分析结果页面截图
    add_image_with_caption(doc, os.path.join(pic_dir, '09_results_page.png'),
                          '分析结果页面 - 展示已保存的分析结果', width=6.5)

    add_paragraph_zh(doc, '在分析结果页面可以：')
    result_features = [
        '查看已保存的所有分析结果',
        '点击结果卡片查看详情',
        '删除不需要的结果',
        '对比不同分析结果'
    ]
    add_bullet_list(doc, result_features)

    doc.add_page_break()

    # ========== 第七章：常见问题与解决方案 ==========
    add_heading_zh(doc, '7. 常见问题与解决方案', level=1)

    add_heading_zh(doc, '7.1 文件上传问题', level=2)

    add_heading_zh(doc, 'Q1: 文件上传失败', level=3)
    add_paragraph_zh(doc, '可能原因及解决方案：')
    q1_solutions = [
        '检查文件格式是否为 CSV/TXT/FCS',
        '检查文件大小（建议 < 500MB）',
        '检查文件编码（推荐 UTF-8）',
        '确认文件未被其他程序占用'
    ]
    add_bullet_list(doc, q1_solutions)

    add_heading_zh(doc, 'Q2: 文件格式不被识别', level=3)
    add_paragraph_zh(doc, '解决方案：')
    q2_solutions = [
        '确保 CSV 文件使用逗号作为分隔符',
        '检查文件第一行是否包含列名',
        '确保数值列使用数字格式而非文本格式'
    ]
    add_bullet_list(doc, q2_solutions)

    add_heading_zh(doc, '7.2 散点图显示问题', level=2)

    add_heading_zh(doc, 'Q3: 散点图显示为一条线或一个点', level=3)
    add_paragraph_zh(doc, '可能原因及解决方案：')
    q3_solutions = [
        '检查选择的维度是否包含数值类型数据',
        '检查数据是否存在异常值',
        '尝试更换降维参数',
        '确认是否选择了正确的列'
    ]
    add_bullet_list(doc, q3_solutions)

    add_heading_zh(doc, 'Q4: UMAP 降维速度慢', level=3)
    add_paragraph_zh(doc, '解决方案：')
    q4_solutions = [
        '减少采样行数（如设置为 1000）',
        '减少选择的维度数量',
        '使用更高配置的机器',
        '关闭其他占用内存的程序'
    ]
    add_bullet_list(doc, q4_solutions)

    add_heading_zh(doc, '7.3 聚类分析问题', level=2)

    add_heading_zh(doc, 'Q5: 聚类分析失败', level=3)
    add_paragraph_zh(doc, '可能原因及解决方案：')
    q5_solutions = [
        '确保已生成散点图',
        '检查网络连接是否正常',
        '查看浏览器控制台错误信息',
        '尝试调整聚类参数（如减小 k 值）'
    ]
    add_bullet_list(doc, q5_solutions)

    add_heading_zh(doc, '7.4 连接问题', level=2)

    add_heading_zh(doc, 'Q6: 前端无法连接后端', level=3)
    add_paragraph_zh(doc, '解决方案：')
    q6_solutions = [
        '检查后端服务是否启动：curl http://localhost:8000/',
        '检查前端 API 配置是否正确',
        '检查防火墙设置',
        '确认端口号未被占用'
    ]
    add_bullet_list(doc, q6_solutions)

    doc.add_page_break()

    # ========== 第八章：附录 ==========
    add_heading_zh(doc, '8. 附录', level=1)

    add_heading_zh(doc, '8.1 环境要求', level=2)

    add_heading_zh(doc, '8.1.1 硬件要求', level=3)
    headers = ['配置项', '最低配置', '推荐配置']
    rows = [
        ['CPU', '4核', '8核及以上'],
        ['内存', '8GB', '16GB及以上'],
        ['硬盘', '10GB可用空间', '50GB及以上'],
        ['网络', '局域网连接', '千兆以太网']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '8.1.2 软件环境', level=3)
    add_paragraph_zh(doc, '前端运行环境：')
    frontend_env = [
        'Node.js 18.x 或更高版本',
        'npm 9.x 或 yarn 1.22.x',
        '现代浏览器（Chrome 90+、Firefox 90+、Edge 90+）'
    ]
    add_bullet_list(doc, frontend_env)

    add_paragraph_zh(doc, '后端运行环境：')
    backend_env = [
        'Python 3.9 或更高版本',
        'pip 21.x 或更高版本'
    ]
    add_bullet_list(doc, backend_env)

    add_heading_zh(doc, '8.2 算法说明', level=2)

    add_heading_zh(doc, '8.2.1 UMAP 算法', level=3)
    add_paragraph_zh(doc,
        'UMAP（Uniform Manifold Approximation and Projection）是一种非线性降维技术，'
        '特别适用于高维单细胞数据的可视化。算法通过构建高维空间的模糊拓扑表示，'
        '优化低维嵌入以保持拓扑结构。')

    add_heading_zh(doc, '8.2.2 Leiden 算法', level=3)
    add_paragraph_zh(doc,
        'Leiden 算法是一种基于图的聚类算法，是 Louvain 算法的改进版本。'
        '算法通过迭代优化模块度来发现社区结构，能够产生更高质量的聚类结果。')

    add_heading_zh(doc, '8.3 性能优化建议', level=2)
    headers = ['数据规模', '建议']
    rows = [
        ['< 10,000 细胞', '直接使用全部数据'],
        ['10,000 - 100,000 细胞', '采样 5,000 - 10,000 细胞'],
        ['> 100,000 细胞', '采样 10,000 细胞或分批处理']
    ]
    add_table_zh(doc, headers, rows)

    add_heading_zh(doc, '8.4 技术支持', level=2)
    add_paragraph_zh(doc, '如有问题或建议，请联系技术支持团队。')
    add_paragraph_zh(doc, '文档版本：v1.0.0')
    add_paragraph_zh(doc, '最后更新：2024年1月')

    # 保存文档
    output_path = '/Users/yangxiaoyu/学习/all/all-projectR/doc/单细胞数据分析系统-用户功能手册（图文版）.docx'
    doc.save(output_path)
    print(f'用户手册已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    create_manual()
