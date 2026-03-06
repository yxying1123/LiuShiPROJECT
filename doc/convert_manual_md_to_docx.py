#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 用户手册转换为 DOCX（基于模板）。
支持：标题、编号列表、无序列表、普通段落、图片。
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path("/Users/yangxiaoyu/学习/all/all-projectR")
INPUT_MD = ROOT / "doc" / "系统功能用户手册-图文版.md"
TEMPLATE_DOCX = ROOT / "doc" / "用户手册模版.docx"
OUTPUT_DOCX = ROOT / "doc" / "系统功能用户手册-图文版.docx"


def set_run_font(run, font_name="SimSun", size=10.5, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_doc_content(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_heading(doc, text, level):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    size_map = {1: 18, 2: 16, 3: 14, 4: 12}
    set_run_font(run, font_name="SimHei", size=size_map.get(level, 12), bold=True)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="SimSun", size=10.5, bold=False)


def add_number_item(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="SimSun", size=10.5, bold=False)


def add_bullet_item(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    set_run_font(run, font_name="SimSun", size=10.5, bold=False)


def add_image(doc, image_raw_path, caption):
    image_path = Path(image_raw_path)
    if not image_path.is_absolute():
        image_path = ROOT / image_path

    if not image_path.exists():
        add_paragraph(doc, f"[图片缺失] {image_raw_path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"图：{caption}")
        set_run_font(cap_run, font_name="SimSun", size=9, bold=False)


def convert(md_path: Path, output_docx: Path, template_docx: Path):
    if not md_path.exists():
        raise FileNotFoundError(f"未找到输入文件: {md_path}")
    if not template_docx.exists():
        raise FileNotFoundError(f"未找到模板文件: {template_docx}")

    doc = Document(str(template_docx))
    clear_doc_content(doc)

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_code_block = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            add_paragraph(doc, line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if re.fullmatch(r"-{3,}", stripped):
            continue

        m_img = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if m_img:
            caption = m_img.group(1).strip()
            img_path = m_img.group(2).strip()
            add_image(doc, img_path, caption)
            continue

        m_h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            add_heading(doc, m_h.group(2).strip(), level)
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            add_number_item(doc, f"{m_num.group(1)}. {m_num.group(2).strip()}")
            continue

        m_bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if m_bullet:
            add_bullet_item(doc, m_bullet.group(1).strip())
            continue

        add_paragraph(doc, stripped)

    doc.save(str(output_docx))


if __name__ == "__main__":
    convert(INPUT_MD, OUTPUT_DOCX, TEMPLATE_DOCX)
    print(f"已生成: {OUTPUT_DOCX}")
